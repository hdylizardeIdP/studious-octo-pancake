import docx
import pdfplumber
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from typing import Optional
import logging

# Configure logging
logger = logging.getLogger(__name__)


class DocumentParser:
    """Handles parsing of various document formats"""

    def parse_text(self, file_path: str) -> str:
        """Parse plain text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def parse_docx(self, file_path: str) -> str:
        """Parse Microsoft Word document"""
        doc = docx.Document(file_path)
        text = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text.strip())

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text.strip())

        return '\n'.join(text)

    def parse_pdf(self, file_path: str) -> str:
        """Parse PDF document"""
        text = []

        # Try pdfplumber first (better for text-based PDFs)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)

            if text:
                return '\n'.join(text)
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyMuPDF")

        # Fallback to PyMuPDF
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text.append(page.get_text())
            doc.close()

            return '\n'.join(text)
        except Exception as e:
            logger.error(f"PyMuPDF also failed: {e}")
            return ""

    def parse_image(self, file_path: str) -> str:
        """Parse image using OCR (Tesseract)"""
        try:
            image = Image.open(file_path)

            # Perform OCR
            text = pytesseract.image_to_string(image)

            return text.strip()
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            # Note: Tesseract needs to be installed on the system
            # For production, consider using AWS Textract or Google Vision API
            raise Exception("OCR processing failed. Ensure Tesseract is installed.")
